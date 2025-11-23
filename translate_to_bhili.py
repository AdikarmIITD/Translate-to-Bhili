import os
import re
import docx
import pdfplumber
import requests
import urllib3
import argparse
from docx import Document

# ---------------- ARGUMENT PARSER ----------------
parser = argparse.ArgumentParser(description="Translate DOCX/PDF to Bhili (DOCX output)")
parser.add_argument('--file', required=True, help="Path to input .docx or .pdf file")
parser.add_argument('--lang', required=True, choices=['hi', 'en'], help="Source language: hi or en")
args = parser.parse_args()

doc = args.file
lang = args.lang

# ---------------- FILE TYPE CHECK ----------------
if doc.lower().endswith('.docx'):
    type_doc = 1
    print("📄 DOCX detected")

elif doc.lower().endswith('.pdf'):
    type_doc = 2
    print("📄 PDF detected")

else:
    raise ValueError("Only .docx and .pdf files are supported.")


# ---------------- CONFIG ----------------
SEPARATOR = '\n'
API_URL = "https://aadivaani.tribal.gov.in/api/translation/translate"
USER_ID = "12"

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

if lang == 'en':
    src_lang = 'eng'
elif lang == 'hi':
    src_lang = 'hin'
else:
    raise ValueError("Language not supported")

target_lang = 'bhili'

print(f"Source language: {src_lang}")
print(f"Target language: {target_lang}")


# ---------------- PDF BLOCK EXTRACTION ----------------
def extract_text_blocks_from_pdf(doc, separator="\n"):
    blocks = []
    with pdfplumber.open(doc) as pdf:
        for page in pdf.pages:
            page_blocks = page.extract_text()
            if page_blocks:
                lines = page_blocks.split("\n")
                for line in lines:
                    line = line.strip()
                    if line:
                        blocks.append(line)
    return separator.join(blocks)


# ---------------- DOCX BLOCK EXTRACTION ----------------
def extract_text_from_docx_corrected(doc_path, separator):
    doc = docx.Document(doc_path)
    extracted_texts = []

    def flatten_text(text):
        return text.replace('\n', ' ').strip()

    for para in doc.paragraphs:
        processed_text = flatten_text(para.text)
        if processed_text:
            extracted_texts.append(processed_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                processed_text = flatten_text(cell.text)
                if processed_text:
                    extracted_texts.append(processed_text)

    return separator.join(extracted_texts)


# ---------------- SENTENCE SPLITTING ----------------
def split_into_sentences(text, lang):
    if lang == "en":
        parts = re.split(r'(?<=[\.\?\!])\s+', text)

    elif lang == "hi":
        parts = re.split(r'(?<=[।\|\?\!])\s*', text)

    else:
        parts = [text]

    return [p.strip() for p in parts if p.strip()]


# ---------------- TRANSLATION ----------------
def translate_adivani(text, source_lang, target_lang, user_id=USER_ID):
    payload = {
        "source_language": source_lang,
        "target_language": target_lang,
        "text": text,
        "user_id": user_id
    }

    try:
        response = requests.post(API_URL, json=payload, verify=False, timeout=15)
        if response.status_code == 200:
            return response.json().get("translated_text", response.text)
        else:
            print(f"Error {response.status_code}: {response.text}")
            return text
    except Exception as e:
        print("Error:", e)
        return text


def translate_block_sentencewise(block, lang, source_lang, target_lang):
    sentences = split_into_sentences(block, lang)
    translated = []

    for s in sentences:
        translated.append(translate_adivani(s, source_lang, target_lang))

    return " ".join(translated)


def translate_file_adivani(input_path, output_path, separator='\n'):
    with open(input_path, 'r', encoding='utf-8') as f:
        blocks = f.read().split(separator)

    translated_blocks = []

    for i, block in enumerate(blocks, 1):
        text = block.strip()

        if not text:
            translated_blocks.append("")
            continue

        print(f"Translating block {i}/{len(blocks)} ...")

        translated_block = translate_block_sentencewise(
            block=text,
            lang=lang,
            source_lang=src_lang,
            target_lang=target_lang
        )

        translated_blocks.append(translated_block)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(separator.join(translated_blocks))

    print(f"✅ Bhili translation saved to {output_path}")


# ---------------- DOCX CREATION FROM PDF ----------------
def reconstruct_docx_from_pdf_blocks(translated_text_path, output_docx_path, separator="\n"):
    with open(translated_text_path, 'r', encoding='utf-8') as f:
        translated_blocks = f.read().split(separator)

    doc = Document()

    for block in translated_blocks:
        doc.add_paragraph(block)

    doc.save(output_docx_path)
    print(f"✅ DOCX created: {output_docx_path}")


# ---------------- DOCX SAFE REPLACE ----------------
def count_blocks(doc_path):
    doc = docx.Document(doc_path)
    count = 0

    for para in doc.paragraphs:
        if para.text.strip():
            count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    count += 1

    return count


def reconstruct_docx_safely(original_path, translated_text_path, output_path, separator):
    with open(translated_text_path, 'r', encoding='utf-8') as f:
        translated_blocks = f.read().split(separator)

    original_block_count = count_blocks(original_path)
    translated_block_count = len(translated_blocks)

    print(f"Original blocks: {original_block_count}")
    print(f"Translated blocks: {translated_block_count}")

    assert original_block_count == translated_block_count, \
        "❌ Block count mismatch. Reconstruction aborted."

    block_iterator = iter(translated_blocks)
    doc = docx.Document(original_path)

    for para in doc.paragraphs:
        if para.text.strip():
            para.clear()
            para.add_run(next(block_iterator))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = next(block_iterator)

    doc.save(output_path)
    print(f"✅ DOCX saved: {output_path}")


# ---------------- MAIN PIPELINE ----------------
if type_doc == 1:
    text_path = "text.txt"

    corrected_text = extract_text_from_docx_corrected(doc, SEPARATOR)

    with open(text_path, 'w', encoding='utf-8') as f:
        f.write(corrected_text)

    print("✅ Text extracted from DOCX")

    output_text_path = "bhili_text.txt"

    translate_file_adivani(
        input_path=text_path,
        output_path=output_text_path
    )

    final_doc_path = "bhili.docx"

    reconstruct_docx_safely(
        original_path=doc,
        translated_text_path=output_text_path,
        output_path=final_doc_path,
        separator=SEPARATOR
    )

    print("\n✅ DOCX → Bhili DOCX completed")


if type_doc == 2:
    text_path = "text.txt"

    extracted_text = extract_text_blocks_from_pdf(doc, SEPARATOR)

    with open(text_path, 'w', encoding='utf-8') as f:
        f.write(extracted_text)

    print("✅ Text extracted from PDF")

    output_text_path = "bhili_text.txt"

    translate_file_adivani(
        input_path=text_path,
        output_path=output_text_path
    )

    final_doc_path = "bhili.docx"

    reconstruct_docx_from_pdf_blocks(
        output_text_path,
        final_doc_path,
        SEPARATOR
    )

    print("\n✅ PDF → Bhili DOCX completed")
